"""
SourceLens — Document Parser
Parses PDF, DOCX, MD, and TXT files into text content.
"""

import os
import logging
from app.models import FileType

logger = logging.getLogger(__name__)


def parse_document(file_path: str, file_type: FileType) -> dict:
    """
    Parse a document and return its text content.
    Returns: {"text": str, "pages": list[dict], "page_count": int}
    """
    if file_type == FileType.PDF:
        return _parse_pdf(file_path)
    elif file_type == FileType.DOCX:
        return _parse_docx(file_path)
    elif file_type in (FileType.MD, FileType.TXT):
        return _parse_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _parse_pdf(file_path: str) -> dict:
    """Parse PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        pages = []
        full_text_parts = []

        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text("text")
            pages.append({
                "page_number": page_num + 1,
                "text": text,
                "char_count": len(text),
            })
            full_text_parts.append(text)

        doc.close()
        return {
            "text": "\n\n".join(full_text_parts),
            "pages": pages,
            "page_count": len(pages),
        }
    except ImportError:
        logger.warning("PyMuPDF not available, falling back to pypdf")
        return _parse_pdf_fallback(file_path)


def _parse_pdf_fallback(file_path: str) -> dict:
    """Fallback PDF parser using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        full_text_parts = []

        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({
                "page_number": page_num + 1,
                "text": text,
                "char_count": len(text),
            })
            full_text_parts.append(text)

        return {
            "text": "\n\n".join(full_text_parts),
            "pages": pages,
            "page_count": len(pages),
        }
    except ImportError:
        logger.error("Neither PyMuPDF nor pypdf is available")
        raise ImportError("No PDF parser available. Install PyMuPDF or pypdf.")


def _parse_docx(file_path: str) -> dict:
    """Parse DOCX using python-docx."""
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text.strip())

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                paragraphs.append(row_text)

    full_text = "\n\n".join(paragraphs)
    return {
        "text": full_text,
        "pages": [{"page_number": 1, "text": full_text, "char_count": len(full_text)}],
        "page_count": 1,  # DOCX doesn't have true page numbers without rendering
    }


def _parse_text(file_path: str) -> dict:
    """Parse plain text or markdown files."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return {
        "text": text,
        "pages": [{"page_number": 1, "text": text, "char_count": len(text)}],
        "page_count": 1,
    }
